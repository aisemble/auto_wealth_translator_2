"""
Document processing module for AutoWealthTranslate.

This module is responsible for parsing PDF and DOCX documents,
extracting text, tables, charts, and formatting information.
"""

import os
import fitz  # PyMuPDF
import docx
import pdfplumber
import pytesseract
import cv2
import numpy as np
from typing import List, Dict, Any, Tuple, Optional, Union
from pathlib import Path
import logging
import re
from dataclasses import dataclass, field

from auto_wealth_translate.utils.logger import get_logger

logger = get_logger(__name__)

@dataclass
class DocumentComponent:
    """Base class for document components."""
    component_id: str
    component_type: str
    page_number: int
    
@dataclass
class TextComponent(DocumentComponent):
    """Text component from a document."""
    text: str
    font_info: Dict[str, Any] = field(default_factory=dict)
    position: Dict[str, float] = field(default_factory=dict)
    is_header: bool = False
    is_footer: bool = False
    
@dataclass
class TableComponent(DocumentComponent):
    """Table component from a document."""
    rows: List[List[str]]
    position: Dict[str, float] = field(default_factory=dict)
    
@dataclass
class ImageComponent(DocumentComponent):
    """Image component from a document."""
    image_data: bytes
    image_format: str
    size: Tuple[int, int]
    position: Dict[str, float] = field(default_factory=dict)
    
@dataclass
class ChartComponent(DocumentComponent):
    """Chart component from a document."""
    chart_data: Dict[str, Any]
    image_data: Optional[bytes] = None
    chart_type: str = "unknown"
    position: Dict[str, float] = field(default_factory=dict)

class DocumentProcessor:
    """
    Process PDF and DOCX documents, extracting components.
    """
    
    def __init__(self, input_file: str):
        """
        Initialize the document processor.
        
        Args:
            input_file: Path to the input document file (PDF or DOCX)
        """
        self.input_file = input_file
        self.file_ext = os.path.splitext(input_file)[1].lower()
        
        if self.file_ext not in ['.pdf', '.docx']:
            raise ValueError(f"Unsupported file format: {self.file_ext}. Supported formats: PDF, DOCX")
            
        self.doc = None
        self.components = []
        logger.info(f"Initialized document processor for {input_file}")
        
    def process(self) -> List[Union[TextComponent, TableComponent, ImageComponent, ChartComponent]]:
        """
        Process the document and extract components.
        
        Returns:
            List of document components (text, tables, images, charts)
        """
        logger.info(f"Processing document: {self.input_file}")
        
        if self.file_ext == '.pdf':
            return self._process_pdf()
        elif self.file_ext == '.docx':
            return self._process_docx()
            
    def _process_pdf(self) -> List[Union[TextComponent, TableComponent, ImageComponent, ChartComponent]]:
        """
        Process a PDF document.
        
        Returns:
            List of document components
        """
        logger.info("Processing PDF document")
        components = []
        component_id = 0
        
        # Use PyMuPDF for text and images
        doc = fitz.open(self.input_file)
        
        # Use pdfplumber for tables
        plumber_pdf = pdfplumber.open(self.input_file)
        
        # First, check if the PDF is scanned (mostly images) and needs OCR
        needs_ocr = self._check_if_needs_ocr(doc)
        
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            plumber_page = plumber_pdf.pages[page_idx]
            
            # Process text using appropriate method
            if needs_ocr:
                logger.info(f"Using OCR for page {page_idx+1} text extraction")
                text_components = self._extract_text_with_ocr(page, page_idx, component_id)
                components.extend(text_components)
                component_id += len(text_components)
            else:
                # Try multiple text extraction methods for better coverage
                text_blocks = page.get_text("blocks")
                if not text_blocks:
                    # If blocks method returns no text, try other methods
                    logger.info(f"No text blocks found on page {page_idx+1}, trying alternate extraction")
                    raw_text = page.get_text("text")
                    if raw_text.strip():
                        # Create a single text component with the entire page text
                        component = TextComponent(
                            component_id=f"text_{component_id}",
                            component_type="text",
                            page_number=page_idx + 1,
                            text=raw_text,
                            font_info={"size": 11},  # Default font size
                            position={
                                "x0": 50,
                                "y0": 50,
                                "x1": page.rect.width - 50,
                                "y1": page.rect.height - 50
                            }
                        )
                        components.append(component)
                        component_id += 1
                else:
                    # Process text blocks
                    for block in text_blocks:
                        # Skip image blocks
                        if block[6] == 1:  # Image block type
                            continue
                            
                        # Create text component
                        text = block[4]
                        if text.strip():
                            # Check if header or footer based on position
                            is_header = block[1] < 100  # Example threshold for header
                            is_footer = block[3] > page.rect.height - 100  # Example threshold for footer
                            
                            component = TextComponent(
                                component_id=f"text_{component_id}",
                                component_type="text",
                                page_number=page_idx + 1,
                                text=text,
                                font_info={"size": block[5]},
                                position={
                                    "x0": block[0],
                                    "y0": block[1],
                                    "x1": block[2],
                                    "y1": block[3]
                                },
                                is_header=is_header,
                                is_footer=is_footer
                            )
                            components.append(component)
                            component_id += 1
            
            # Extract tables
            tables = plumber_page.find_tables()
            for table in tables:
                rows = []
                for row in table.extract():
                    rows.append([str(cell) if cell is not None else "" for cell in row])
                
                component = TableComponent(
                    component_id=f"table_{component_id}",
                    component_type="table",
                    page_number=page_idx + 1,
                    rows=rows,
                    position={
                        "x0": table.bbox[0],
                        "y0": table.bbox[1],
                        "x1": table.bbox[2],
                        "y1": table.bbox[3]
                    }
                )
                components.append(component)
                component_id += 1
            
            # Extract images
            image_list = page.get_images(full=True)
            for img_idx, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Simplified image rectangle handling - create a default rectangle if unable to get exact position
                img_rect = fitz.Rect(100, 100, 400, 400)  # Default position
                
                # Add image component with position information
                component = ImageComponent(
                    component_id=f"image_{component_id}",
                    component_type="image",
                    page_number=page_idx + 1,
                    image_data=image_bytes,
                    image_format=image_ext,
                    size=(base_image["width"], base_image["height"]),
                    position={
                        "x0": img_rect.x0,
                        "y0": img_rect.y0,
                        "x1": img_rect.x1,
                        "y1": img_rect.y1
                    }
                )
                components.append(component)
                component_id += 1
        
        # If we didn't extract any text, try a more aggressive approach
        if not any(isinstance(c, TextComponent) for c in components):
            logger.warning("No text components extracted, trying alternative method")
            # Try to extract text from the whole document using a different approach
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                text = page.get_text("text")
                if text.strip():
                    component = TextComponent(
                        component_id=f"text_{component_id}",
                        component_type="text",
                        page_number=page_idx + 1,
                        text=text,
                        font_info={"size": 11},
                        position={
                            "x0": 50,
                            "y0": 50,
                            "x1": page.rect.width - 50,
                            "y1": page.rect.height - 50
                        }
                    )
                    components.append(component)
                    component_id += 1
        
        plumber_pdf.close()
        doc.close()
        
        logger.info(f"Extracted {len(components)} components from PDF")
        return components
    
    def _check_if_needs_ocr(self, doc):
        """
        Check if the PDF appears to be scanned and needs OCR.
        
        Args:
            doc: PyMuPDF document
            
        Returns:
            Boolean indicating if OCR is needed
        """
        # Sample a few pages to see if they contain searchable text
        total_pages = len(doc)
        sample_pages = min(total_pages, 3)  # Check up to 3 pages
        
        text_count = 0
        image_count = 0
        
        for page_idx in range(sample_pages):
            page = doc[page_idx]
            text = page.get_text("text")
            images = page.get_images(full=True)
            
            if text.strip():
                text_count += 1
            if images:
                image_count += len(images)
        
        # If there are significantly more images than text pages, might need OCR
        return (text_count == 0 and image_count > 0) or (image_count > text_count * 3)
    
    def _extract_text_with_ocr(self, page, page_idx, start_component_id):
        """
        Extract text from a page using OCR.
        
        Args:
            page: PyMuPDF page
            page_idx: Page index
            start_component_id: Starting component ID
            
        Returns:
            List of extracted TextComponents
        """
        components = []
        component_id = start_component_id
        
        try:
            # Render page to image for OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
            
            # Convert to grayscale if it's not already
            if pix.n > 1:
                gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            else:
                gray = img
            
            # Use Tesseract OCR to extract text
            text = pytesseract.image_to_string(gray, lang='eng')
            
            if text.strip():
                # Create a text component for the OCR text
                component = TextComponent(
                    component_id=f"text_{component_id}",
                    component_type="text",
                    page_number=page_idx + 1,
                    text=text,
                    font_info={"size": 11},  # Default font size for OCR text
                    position={
                        "x0": 50,
                        "y0": 50,
                        "x1": page.rect.width - 50,
                        "y1": page.rect.height - 50
                    },
                    is_header=False,
                    is_footer=False
                )
                components.append(component)
                component_id += 1
        except Exception as e:
            logger.error(f"OCR processing error: {str(e)}")
        
        return components
        
    def _process_docx(self) -> List[Union[TextComponent, TableComponent, ImageComponent, ChartComponent]]:
        """
        Process a DOCX document.
        
        Returns:
            List of document components
        """
        logger.info("Processing DOCX document")
        components = []
        component_id = 0
        
        doc = docx.Document(self.input_file)
        
        # Process paragraphs (text)
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                # Simplified - in a real implementation, you'd extract more paragraph formatting
                component = TextComponent(
                    component_id=f"text_{component_id}",
                    component_type="text",
                    page_number=0,  # DOCX doesn't have direct page mapping
                    text=para.text,
                    font_info={},  # Would extract font info in full implementation
                    position={}     # DOCX doesn't have direct position info
                )
                components.append(component)
                component_id += 1
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                rows.append(row_data)
                
            component = TableComponent(
                component_id=f"table_{component_id}",
                component_type="table",
                page_number=0,  # DOCX doesn't have direct page mapping
                rows=rows,
                position={}     # DOCX doesn't have direct position info
            )
            components.append(component)
            component_id += 1
                
        # Process images (simplified)
        # For a complete implementation, you would extract the actual image data
        
        logger.info(f"Extracted {len(components)} components from DOCX")
        return components
