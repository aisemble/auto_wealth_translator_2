"""
Document rebuilding module for AutoWealthTranslate.

This module is responsible for reconstructing documents with 
translated content while preserving formatting.
"""

import os
import io
import tempfile
from typing import List, Dict, Any, Union, Optional, Tuple
from pathlib import Path
import fitz  # PyMuPDF
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
import logging
import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageColor

from auto_wealth_translate.utils.logger import get_logger
from auto_wealth_translate.core.document_processor import (
    DocumentComponent, TextComponent, TableComponent, 
    ImageComponent, ChartComponent
)

logger = get_logger(__name__)

class DocumentOutput:
    """Class representing a built document."""
    
    def __init__(self, data: bytes, format: str):
        """
        Initialize document output.
        
        Args:
            data: Document data as bytes
            format: Document format (pdf, docx)
        """
        self.data = data
        self.format = format
        
    def save(self, output_path: str) -> None:
        """
        Save the document to a file.
        
        Args:
            output_path: Path to save the document to
        """
        with open(output_path, 'wb') as f:
            f.write(self.data)
        logger.info(f"Document saved to {output_path}")

class DocumentRebuilder:
    """
    Rebuilds documents with translated content.
    """
    
    # Rebuild modes
    MODE_ENHANCED = "enhanced"      # Enhanced version of our current approach
    MODE_PRECISE = "precise"        # Precise element identification (Canva-like)
    MODE_BILINGUAL = "bilingual"    # Side-by-side or sequential bilingual pages
    MODE_BILINGUAL_MARKDOWN = "bilingual_markdown"  # Bilingual mode with markdown translation
    
    def __init__(self):
        """Initialize the document rebuilder."""
        logger.info("Initialized document rebuilder")
        self.font_cache = {}  # Cache fonts to avoid reloading them
        
    def rebuild(self, components: List[DocumentComponent], output_format: str = 'pdf', 
                rebuild_mode: str = MODE_ENHANCED, source_pdf_path: str = None,
                source_lang: str = None, target_lang: str = None, translation_model: str = None) -> DocumentOutput:
        """
        Rebuild a document from components.
        
        Args:
            components: List of document components
            output_format: Format of the output document (pdf, docx)
            rebuild_mode: Mode for rebuilding the document 
                          ('enhanced', 'precise', 'bilingual', or 'bilingual_markdown')
            source_pdf_path: Path to original PDF (needed for certain modes)
            source_lang: Source language code (needed for bilingual_markdown mode)
            target_lang: Target language code (needed for bilingual_markdown mode)
            translation_model: Translation model to use (needed for bilingual_markdown mode)
            
        Returns:
            DocumentOutput object
        """
        logger.info(f"Rebuilding document in {output_format} format using {rebuild_mode} mode")
        
        # Sort components by page number and position
        sorted_components = sorted(
            components, 
            key=lambda c: (c.page_number, getattr(c, 'position', {}).get('y0', 0))
        )
        
        if output_format.lower() == 'pdf':
            if rebuild_mode == self.MODE_ENHANCED:
                return self._rebuild_pdf_enhanced(sorted_components)
            elif rebuild_mode == self.MODE_PRECISE:
                if not source_pdf_path:
                    logger.warning("Source PDF path required for precise mode. Falling back to enhanced mode.")
                    return self._rebuild_pdf_enhanced(sorted_components)
                return self._rebuild_pdf_precise(sorted_components, source_pdf_path)
            elif rebuild_mode == self.MODE_BILINGUAL:
                if not source_pdf_path:
                    logger.warning("Source PDF path required for bilingual mode. Falling back to enhanced mode.")
                    return self._rebuild_pdf_enhanced(sorted_components)
                return self._rebuild_pdf_bilingual(sorted_components, source_pdf_path)
            elif rebuild_mode == self.MODE_BILINGUAL_MARKDOWN:
                if not source_pdf_path:
                    logger.warning("Source PDF path required for bilingual markdown mode. Falling back to enhanced mode.")
                    return self._rebuild_pdf_enhanced(sorted_components)
                if not all([source_lang, target_lang, translation_model]):
                    logger.warning("Language parameters required for bilingual markdown mode. Falling back to standard bilingual mode.")
                    return self._rebuild_pdf_bilingual(sorted_components, source_pdf_path)
                return self._rebuild_pdf_bilingual_markdown(sorted_components, source_pdf_path, 
                                                          source_lang, target_lang, translation_model)
            else:
                logger.warning(f"Unknown rebuild mode: {rebuild_mode}. Using enhanced mode.")
                return self._rebuild_pdf_enhanced(sorted_components)
        elif output_format.lower() == 'docx':
            return self._rebuild_docx(sorted_components)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _find_cjk_font(self):
        """Find a suitable CJK font on the system."""
        # Check common font directories by platform
        system_font_dirs = []
        if os.name == 'posix':  # macOS, Linux
            system_font_dirs = [
                "/System/Library/Fonts",  # macOS
                "/Library/Fonts",         # macOS
                "/usr/share/fonts",       # Linux
                "/usr/local/share/fonts"  # Linux
            ]
        elif os.name == 'nt':  # Windows
            system_font_dirs = [
                "C:\\Windows\\Fonts"
            ]
        
        # List of CJK font filenames to look for
        cjk_font_names = [
            # MacOS
            "PingFang.ttc", "PingFangSC-Regular.ttf", "STHeiti Light.ttc", "Hiragino Sans GB.ttc",
            # Windows
            "msyh.ttf", "simhei.ttf", "simsun.ttc", "simkai.ttf",
            # Linux
            "NotoSansCJK-Regular.ttc", "NotoSansSC-Regular.otf", "WenQuanYiMicroHei.ttf"
        ]
        
        # Try to find a suitable font
        for font_dir in system_font_dirs:
            if os.path.exists(font_dir):
                for font_name in cjk_font_names:
                    font_path = os.path.join(font_dir, font_name)
                    if os.path.exists(font_path):
                        logger.info(f"Found CJK font: {font_path}")
                        return font_path
        
        logger.warning("No CJK font found in system. Chinese characters may not display correctly.")
        return None
    
    def _get_font(self, font_path: str, size: int) -> ImageFont.FreeTypeFont:
        """Get a font from cache or load it."""
        cache_key = f"{font_path}_{size}"
        if cache_key not in self.font_cache:
            try:
                self.font_cache[cache_key] = ImageFont.truetype(font_path, size)
            except Exception as e:
                logger.error(f"Error loading font {font_path}: {str(e)}")
                # Fall back to default font
                self.font_cache[cache_key] = ImageFont.load_default()
        return self.font_cache[cache_key]
    
    def _get_text_dimensions(self, text: str, font: ImageFont.FreeTypeFont) -> Tuple[int, int]:
        """Get the width and height of text with a given font."""
        # PIL 9.5.0 uses getsize(), but newer versions use getbbox()
        if hasattr(font, 'getsize'):
            return font.getsize(text)
        elif hasattr(font, 'getbbox'):
            # Convert bounding box to dimensions
            left, top, right, bottom = font.getbbox(text)
            return (right - left, bottom - top)
        else:
            # Rough estimate if we can't determine
            return (len(text) * font.size // 2, font.size)

    def _get_text_by_component_id(self, components, component_id):
        """Get text from the component list by component_id."""
        for component in components:
            if component.component_id == component_id:
                if isinstance(component, TextComponent):
                    return component.text
        return None
            
    def _rebuild_pdf_enhanced(self, components: List[DocumentComponent]) -> DocumentOutput:
        """
        Rebuild a PDF document with improved layout preservation (enhanced mode).
        This is an improvement of our previous approach.
        
        Args:
            components: List of document components
            
        Returns:
            DocumentOutput object with PDF data
        """
        # Create a temporary file for the output PDF
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_file.close()
        output_path = temp_file.name
        
        # Find a suitable CJK font
        cjk_font_path = self._find_cjk_font()
        
        # Fallback font paths for different systems
        fallback_fonts = []
        if os.name == 'posix':  # macOS, Linux
            fallback_fonts = [
                "/System/Library/Fonts/STHeiti Light.ttc",  # macOS
                "/System/Library/Fonts/PingFang.ttc",       # macOS
                "/usr/share/fonts/truetype/arphic/uming.ttc", # Linux
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc" # Linux
            ]
        elif os.name == 'nt':  # Windows
            fallback_fonts = [
                "C:\\Windows\\Fonts\\msyh.ttf",
                "C:\\Windows\\Fonts\\simsun.ttc"
            ]
        
        # Try loading fallback fonts if no CJK font was found
        if not cjk_font_path:
            for font_path in fallback_fonts:
                if os.path.exists(font_path):
                    cjk_font_path = font_path
                    logger.info(f"Using fallback CJK font: {font_path}")
                    break
        
        # If still no CJK font, use a built-in PDF font as last resort
        default_font_path = cjk_font_path if cjk_font_path else "helv"
        
        # Group components by page
        page_components = {}
        max_page = 0
        for component in components:
            page_num = component.page_number
            max_page = max(max_page, page_num)
            if page_num not in page_components:
                page_components[page_num] = []
            page_components[page_num].append(component)
        
        # Create a PDF document
        doc = fitz.open()
        
        # Constants for image resolution
        DPI = 300  # Higher DPI for better quality
        SCALE = DPI / 72  # Scale factor (72 is the default PDF DPI)
        
        # For each page
        for page_num in range(1, max_page + 1):
            if page_num not in page_components:
                continue
                
            # Create a new page (Letter size or A4)
            page_width, page_height = 612, 792  # US Letter in points
            page = doc.new_page(width=page_width, height=page_height)
            
            # Create an image of the page for PIL-based text rendering
            img_width, img_height = int(page_width * SCALE), int(page_height * SCALE)
            img = Image.new('RGB', (img_width, img_height), (255, 255, 255))
            draw = ImageDraw.Draw(img)
            
            # Sort components by vertical position for better layout
            page_components[page_num].sort(key=lambda c: getattr(c, 'position', {}).get('y0', 0))
            
            # First handle images and background elements
            for component in page_components[page_num]:
                if isinstance(component, ImageComponent) and component.image_data:
                    position = component.position
                    if not position:
                        position = {"x0": 50, "y0": 300, "x1": 550, "y1": 550}
                    
                    rect = fitz.Rect(
                        position.get('x0', 50),
                        position.get('y0', 50),
                        position.get('x1', 550),
                        position.get('y1', 550)
                    )
                    
                    try:
                        page.insert_image(rect, stream=component.image_data)
                    except Exception as e:
                        logger.error(f"Error adding image: {str(e)}")
            
            # Then handle text and tables - they will be drawn on top of images
            for component in page_components[page_num]:
                if isinstance(component, TextComponent):
                    text = component.text
                    if not text.strip():
                        continue
                    
                    # Get position info or use defaults
                    position = component.position
                    if not position:
                        position = {"x0": 50, "y0": 50, "x1": 550, "y1": 70}
                    
                    # Scale position to image size
                    x0 = position.get('x0', 50) * SCALE
                    y0 = position.get('y0', 50) * SCALE
                    x1 = position.get('x1', 550) * SCALE
                    y1 = position.get('y1', 70) * SCALE
                    
                    # Font size calculation - more precise than before
                    original_font_size = component.font_info.get('size', 11)
                    
                    # Scale font size to match image resolution
                    font_size = int(max(original_font_size * SCALE, 16))  # Minimum readable size
                    
                    # Try to use PIL to draw text with CJK support
                    try:
                        # Determine if text contains non-Latin characters
                        has_non_latin = any(ord(char) > 255 for char in text)
                        
                        # Choose appropriate font
                        font_path = cjk_font_path if (has_non_latin and cjk_font_path) else default_font_path
                        
                        # Use font from path if it's a file path, otherwise use the name directly with PyMuPDF
                        if os.path.exists(str(font_path)):
                            # Get or load font for PIL
                            font = self._get_font(font_path, font_size)
                            
                            if has_non_latin and cjk_font_path:
                                logger.info(f"Using CJK font for text with non-Latin characters")
                            
                            # Calculate text dimensions for positioning
                            text_width, text_height = self._get_text_dimensions(text, font)
                            
                            # Determine text color (black by default)
                            text_color = (0, 0, 0)
                            
                            # Check for special text formatting
                            is_bold = component.font_info.get('bold', False)
                            is_italic = component.font_info.get('italic', False)
                            
                            # Use different color for headers
                            if component.is_header:
                                text_color = (0, 0, 180)  # Dark blue for headers
                            
                            # Align text based on available space
                            text_y = y0 + (y1 - y0 - text_height) / 2  # Center vertically
                            
                            # Check if the text would overflow and needs wrapping
                            available_width = x1 - x0
                            if text_width > available_width:
                                # Text needs wrapping
                                words = text.split()
                                lines = []
                                current_line = []
                                current_width = 0
                                
                                for word in words:
                                    word_width, _ = self._get_text_dimensions(word + " ", font)
                                    if current_width + word_width <= available_width:
                                        current_line.append(word)
                                        current_width += word_width
                                    else:
                                        if current_line:
                                            lines.append(" ".join(current_line))
                                        current_line = [word]
                                        current_width = word_width
                                
                                if current_line:
                                    lines.append(" ".join(current_line))
                                
                                # Draw each line of text
                                line_spacing = font_size * 1.2  # 120% of font size
                                for i, line in enumerate(lines):
                                    line_y = text_y + i * line_spacing
                                    if line_y + text_height < img_height:  # Ensure line is within page
                                        draw.text((x0, line_y), line, font=font, fill=text_color)
                            else:
                                # Draw single line text
                                draw.text((x0, text_y), text, font=font, fill=text_color)
                        else:
                            # We have a font name instead of a path, use PyMuPDF directly
                            # This will render text directly on the page instead of using PIL
                            # Scale back from image size to PDF coordinates
                            pdf_rect = fitz.Rect(
                                x0 / SCALE, 
                                y0 / SCALE, 
                                x1 / SCALE, 
                                y1 / SCALE
                            )
                            
                            font_name = "china" if has_non_latin else "helv"
                            text_color_pdf = (0, 0, 0.7) if component.is_header else (0, 0, 0)
                            
                            # Insert text directly into PDF
                            page.insert_textbox(
                                pdf_rect,
                                text,
                                fontname=font_name,
                                fontsize=original_font_size,
                                color=text_color_pdf,
                                align=fitz.TEXT_ALIGN_LEFT
                            )
                        
                    except Exception as e:
                        logger.error(f"Error drawing text with PIL: {str(e)}")
                        # Try PyMuPDF as fallback
                        try:
                            pdf_rect = fitz.Rect(
                                position.get('x0', 50),
                                position.get('y0', 50),
                                position.get('x1', 550),
                                position.get('y1', 70)
                            )
                            
                            font_name = "china" if has_non_latin else "helv"
                            
                            page.insert_textbox(
                                pdf_rect,
                                text,
                                fontname=font_name,
                                fontsize=original_font_size,
                                align=fitz.TEXT_ALIGN_LEFT
                            )
                        except Exception as e2:
                            logger.error(f"Fallback text insertion also failed: {str(e2)}")
                
                elif isinstance(component, TableComponent):
                    # Get position info
                    position = component.position
                    if not position:
                        position = {"x0": 50, "y0": 200, "x1": 550, "y1": 400}
                    
                    # Scale position to image size
                    x0 = position.get('x0', 50) * SCALE
                    y0 = position.get('y0', 50) * SCALE
                    x1 = position.get('x1', 550) * SCALE
                    y1 = position.get('y1', 400) * SCALE
                    
                    # Draw table
                    self._draw_table_on_image(draw, component, (x0, y0, x1, y1), cjk_font_path, default_font_path)
            
            # Check if we used PIL for rendering - if so, add the image to the PDF
            if isinstance(default_font_path, str) and os.path.exists(default_font_path):
                # Resize and optimize the PIL image
                img_resized = img.resize((page_width, page_height), Image.Resampling.LANCZOS)
                
                # Convert to bytes
                img_bytes = io.BytesIO()
                img_resized.save(img_bytes, format='PNG', optimize=True)
                img_bytes.seek(0)
                
                # Add the PIL-rendered image to the PDF page
                try:
                    # Insert the rendered image as the page background
                    page.insert_image(fitz.Rect(0, 0, page_width, page_height), stream=img_bytes.getvalue())
                except Exception as e:
                    logger.error(f"Error adding rendered page to PDF: {str(e)}")
        
        # Save the document
        doc.save(output_path)
        doc.close()
        
        # Read the file back
        with open(output_path, 'rb') as f:
            pdf_data = f.read()
        
        # Clean up the temporary file
        try:
            os.unlink(output_path)
        except:
            pass
        
        return DocumentOutput(pdf_data, 'pdf')
    
    def _rebuild_pdf_precise(self, components: List[DocumentComponent], source_pdf_path: str) -> DocumentOutput:
        """
        Rebuild a PDF document using precise object identification (Canva-like approach).
        This method preserves the original PDF layout exactly and only replaces text.
        
        Args:
            components: List of document components
            source_pdf_path: Path to the original PDF file
            
        Returns:
            DocumentOutput object with PDF data
        """
        # Create a temporary file for the output PDF
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_file.close()
        output_path = temp_file.name
        
        try:
            # Open the source PDF
            doc = fitz.open(source_pdf_path)
            
            # Find a suitable CJK font for Chinese characters
            cjk_font_path = self._find_cjk_font()
            
            # Build a dictionary of components by page number
            components_by_page = {}
            for component in components:
                page_num = component.page_number
                if page_num not in components_by_page:
                    components_by_page[page_num] = []
                components_by_page[page_num].append(component)
            
            # Process each page
            for page_num, page in enumerate(doc, start=1):
                if page_num not in components_by_page:
                    continue
                
                # Get page components
                page_components = components_by_page[page_num]
                
                # Create a list of text components
                text_components = [c for c in page_components if isinstance(c, TextComponent)]
                
                # First extract detailed text information using the "dict" mode
                text_dict = page.get_text("dict")
                
                # Process text blocks and spans for replacement
                for block in text_dict["blocks"]:
                    if block["type"] == 0:  # If it's a text block
                        for line in block["lines"]:
                            for span in line["spans"]:
                                span_rect = fitz.Rect(span["bbox"])
                                span_text = span["text"].strip()
                                
                                if not span_text:
                                    continue
                                
                                # Try to find a matching component by position overlap
                                matched_component = None
                                for comp in text_components:
                                    pos = comp.position
                                    if pos:
                                        comp_rect = fitz.Rect(pos.get('x0'), pos.get('y0'), pos.get('x1'), pos.get('y1'))
                                        # Check if the rectangles overlap significantly
                                        overlap = span_rect.intersect(comp_rect)
                                        if overlap.get_area() > 0.5 * min(span_rect.get_area(), comp_rect.get_area()):
                                            matched_component = comp
                                            break
                                
                                # If we found a matching component, replace the text
                                if matched_component:
                                    try:
                                        translated_text = matched_component.text.strip()
                                        if not translated_text:
                                            continue
                                            
                                        # Check if translated text contains non-Latin characters
                                        has_non_latin = any(ord(char) > 255 for char in translated_text)
                                        
                                        # Get original span formatting
                                        orig_font = span.get("font", "helv")
                                        orig_size = span.get("size", 11)
                                        orig_color = span.get("color", (0, 0, 0))
                                        
                                        # Determine the most appropriate font for the content
                                        font_name = "china" if has_non_latin else orig_font
                                        
                                        # Create a redaction annotation to cover the original text
                                        annot = page.add_redact_annot(span_rect, fill=(1, 1, 1))
                                        page.apply_redactions()
                                        
                                        # Handle text size adaptation if needed
                                        # Calculate a rough estimate of width needed
                                        text_width_estimate = len(translated_text) * orig_size * 0.6
                                        if text_width_estimate > span_rect.width * 1.2:
                                            # Text is significantly longer, adjust font size
                                            scale_factor = min(1.0, span_rect.width * 1.2 / text_width_estimate)
                                            adjusted_size = max(orig_size * scale_factor, 8)  # Don't go below 8pt
                                        else:
                                            adjusted_size = orig_size
                                        
                                        # Insert the translated text with appropriate font and formatting
                                        page.insert_textbox(
                                            span_rect,
                                            translated_text,
                                            fontname=font_name,
                                            fontsize=adjusted_size,
                                            color=orig_color,
                                            align=fitz.TEXT_ALIGN_LEFT
                                        )
                                    except Exception as e:
                                        logger.error(f"Error replacing text: {str(e)}")
                
                # Now handle any image components if needed
                # For this version, we'll keep the original images
            
            # Save the document
            doc.save(output_path)
            doc.close()
            
            # Read the file back
            with open(output_path, 'rb') as f:
                pdf_data = f.read()
            
        except Exception as e:
            logger.error(f"Error in precise PDF rebuilding: {str(e)}")
            # Fallback to enhanced mode
            logger.warning("Falling back to enhanced mode")
            return self._rebuild_pdf_enhanced(components)
        
        # Clean up the temporary file
        try:
            os.unlink(output_path)
        except:
            pass
        
        return DocumentOutput(pdf_data, 'pdf')
    
    def _rebuild_pdf_bilingual(self, components: List[DocumentComponent], source_pdf_path: str) -> DocumentOutput:
        """
        Rebuild a PDF document with bilingual pages.
        This method keeps the original pages and adds translated pages after each original.
        
        Args:
            components: List of document components
            source_pdf_path: Path to the original PDF file
            
        Returns:
            DocumentOutput object with PDF data
        """
        # Create a temporary file for the output PDF
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_file.close()
        output_path = temp_file.name
        
        try:
            # Open the source PDF to get original pages
            orig_doc = fitz.open(source_pdf_path)
            
            # Create a new document for the bilingual output
            new_doc = fitz.open()
            
            # Find a suitable CJK font
            cjk_font_path = self._find_cjk_font()
            
            # Group components by page number
            page_components = {}
            max_page = 0
            for component in components:
                page_num = component.page_number
                max_page = max(max_page, page_num)
                if page_num not in page_components:
                    page_components[page_num] = []
                page_components[page_num].append(component)
            
            # For each page in the original document
            for page_idx in range(len(orig_doc)):
                page_num = page_idx + 1
                
                # Add the original page
                new_doc.insert_pdf(orig_doc, from_page=page_idx, to_page=page_idx)
                
                # If we have translated components for this page, add a translation page
                if page_num in page_components:
                    # Create a translation page
                    page = new_doc.new_page(width=orig_doc[page_idx].rect.width, 
                                      height=orig_doc[page_idx].rect.height)
                    
                    # Sort components by position
                    page_components[page_num].sort(key=lambda c: (
                        getattr(c, 'position', {}).get('y0', 0),
                        getattr(c, 'position', {}).get('x0', 0)
                    ))
                    
                    # Add a header to indicate this is a translation page
                    header_rect = fitz.Rect(50, 30, page.rect.width - 50, 60)
                    try:
                        # Use built-in fonts for the header which support unicode
                        page.insert_textbox(
                            header_rect,
                            f"Translation - Page {page_num}",
                            fontsize=14,
                            color=(0, 0, 0.8),  # Dark blue
                            fontname="helv",  # Use a built-in font with Unicode support
                            align=fitz.TEXT_ALIGN_CENTER
                        )
                    except Exception as e:
                        logger.error(f"Error adding header: {str(e)}")
                        # Fallback to simple text insertion
                        try:
                            page.insert_text((header_rect.x0, header_rect.y0 + 14), 
                                             f"Translation - Page {page_num}", 
                                             fontsize=14, 
                                             color=(0, 0, 0.8))
                        except Exception as e2:
                            logger.error(f"Simple header insertion also failed: {str(e2)}")
                    
                    # Process text components
                    y_position = 70
                    y_spacing = 15  # Standard spacing between items
                    
                    for component in page_components[page_num]:
                        # Only process text components in the bilingual mode
                        if isinstance(component, TextComponent):
                            text = component.text
                            
                            # Create a text rectangle for this component
                            text_rect = fitz.Rect(50, y_position, page.rect.width - 50, y_position + 40)
                            
                            # Calculate appropriate font size
                            original_size = component.font_size if hasattr(component, 'font_size') else 11
                            display_font_size = min(original_size, 11)  # Limit to reasonable size
                            
                            # For small text, use a minimum readable size
                            display_font_size = max(display_font_size, 9)
                            
                            # Add the text
                            try:
                                # Check if text contains non-Latin characters
                                has_non_latin = any(ord(char) > 255 for char in text)
                                
                                # Choose font based on content
                                font_name = "china" if has_non_latin else "helv"
                                
                                # Insert text - text box will handle wrapping
                                page.insert_textbox(
                                    text_rect,
                                    text,
                                    fontsize=display_font_size,
                                    fontname=font_name,  # Use built-in font with CJK support
                                    align=fitz.TEXT_ALIGN_LEFT
                                )
                                
                                # Try simpler text handling approach
                                try:
                                    # Try direct text insertion with line breaks
                                    y_pos = text_rect.y0
                                    line_height = display_font_size * 1.2
                                    lines = text.split('\n')
                                    for line in lines:
                                        page.insert_text((text_rect.x0, y_pos), line, 
                                                        fontsize=display_font_size)
                                        y_pos += line_height
                                except Exception as e_simple:
                                    logger.error(f"Even simple text insertion failed: {str(e_simple)}")
                                
                                # FIX: Don't use get_textbox_layout, estimate text height instead
                                # Estimate the number of lines based on text length and available width
                                estimated_text_width = len(text) * display_font_size * 0.6  # Rough estimate
                                available_width = page.rect.width - 100  # 50px padding on each side
                                estimated_lines = max(1, int(estimated_text_width / available_width) + 1)
                                
                                # Add height for all estimated lines
                                y_position += estimated_lines * y_spacing
                                
                            except Exception as e:
                                logger.error(f"Error adding translated text: {str(e)}")
                                y_position += y_spacing
                        
                        # Add support for image components in bilingual mode
                        elif isinstance(component, ImageComponent) and component.image_data:
                            try:
                                # Calculate image size to fit on the page
                                img_width = component.size[0]
                                img_height = component.size[1]
                                
                                # Scale to fit on page width with margins
                                max_width = page.rect.width - 100  # 50px margin on each side
                                scale_factor = min(1.0, max_width / img_width)
                                
                                display_width = img_width * scale_factor
                                display_height = img_height * scale_factor
                                
                                # Center image horizontally
                                x0 = (page.rect.width - display_width) / 2
                                # Position image at current y position
                                rect = fitz.Rect(x0, y_position, x0 + display_width, y_position + display_height)
                                
                                # Add the image
                                page.insert_image(rect, stream=component.image_data)
                                
                                # Update y position for next element
                                y_position += display_height + y_spacing
                            except Exception as e:
                                logger.error(f"Error adding image to bilingual PDF: {str(e)}")
                                y_position += y_spacing
            
            # Save the document
            new_doc.save(output_path)
            orig_doc.close()
            new_doc.close()
            
            # Read the file back
            with open(output_path, 'rb') as f:
                pdf_data = f.read()
                
        except Exception as e:
            logger.error(f"Error in bilingual PDF rebuilding: {str(e)}")
            # Fallback to enhanced mode
            logger.warning("Falling back to enhanced mode")
            return self._rebuild_pdf_enhanced(components)
        
        # Clean up the temporary file
        try:
            os.unlink(output_path)
        except:
            pass
        
        return DocumentOutput(pdf_data, 'pdf')
    
    def _rebuild_pdf_bilingual_markdown(self, components: List[DocumentComponent], source_pdf_path: str, 
                                        source_lang: str, target_lang: str, translation_model: str) -> DocumentOutput:
        """
        Rebuild a PDF document with bilingual pages using the Markdown approach.
        This method keeps the original pages and adds translated pages after each original,
        using the Markdown processor for better translation quality and formatting.
        
        Args:
            components: List of document components
            source_pdf_path: Path to the original PDF file
            source_lang: Source language code
            target_lang: Target language code
            translation_model: Translation model to use
            
        Returns:
            DocumentOutput object with PDF data
        """
        # Import needed modules here to avoid circular imports
        from auto_wealth_translate.core.markdown_processor import MarkdownProcessor
        from auto_wealth_translate.core.translator import TranslationService
        
        # Create a temporary file for the output PDF
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_file.close()
        output_path = temp_file.name
        
        try:
            # Open the source PDF to get original pages
            orig_doc = fitz.open(source_pdf_path)
            
            # Create a new document for the bilingual output
            new_doc = fitz.open()
            
            # Initialize the markdown processor and translator
            md_processor = MarkdownProcessor()
            translation_service = TranslationService(
                source_lang=source_lang,
                target_lang=target_lang,
                model=translation_model
            )
            
            # For each page in the original document
            for page_idx in range(len(orig_doc)):
                page_num = page_idx + 1
                logger.info(f"Processing page {page_num} for bilingual markdown output")
                
                # Add the original page
                new_doc.insert_pdf(orig_doc, from_page=page_idx, to_page=page_idx)
                
                # Create a temporary PDF for just this page
                temp_page_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_page_pdf.close()
                page_pdf_path = temp_page_pdf.name
                
                # Extract just this page to a separate PDF
                with fitz.open() as page_doc:
                    page_doc.insert_pdf(orig_doc, from_page=page_idx, to_page=page_idx)
                    page_doc.save(page_pdf_path)
                
                try:
                    # Convert this page to markdown
                    logger.info(f"Converting page {page_num} to markdown")
                    md_content = md_processor.pdf_to_markdown(page_pdf_path)
                    
                    # Translate the markdown content
                    logger.info(f"Translating page {page_num} content")
                    translated_md = md_processor.translate_markdown(md_content, translation_service)
                    
                    # Create a temporary PDF from the translated markdown
                    temp_translated_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                    temp_translated_pdf.close()
                    translated_pdf_path = temp_translated_pdf.name
                    
                    # Generate PDF from the translated markdown
                    logger.info(f"Generating translated PDF for page {page_num}")
                    md_processor.markdown_to_pdf(translated_md, translated_pdf_path)
                    
                    # Add the translated page(s) from the generated PDF
                    with fitz.open(translated_pdf_path) as translated_doc:
                        # Add header to indicate this is a translation page
                        for t_page in translated_doc:
                            # Create a new page with the same dimensions
                            new_page = new_doc.new_page(width=t_page.rect.width, height=t_page.rect.height)
                            
                            # Add a header
                            header_rect = fitz.Rect(50, 30, new_page.rect.width - 50, 60)
                            try:
                                new_page.insert_textbox(
                                    header_rect,
                                    f"Translation - Page {page_num}",
                                    fontsize=14,
                                    color=(0, 0, 0.8),  # Dark blue
                                    fontname="helv",
                                    align=fitz.TEXT_ALIGN_CENTER
                                )
                            except Exception as header_err:
                                logger.error(f"Error adding header: {str(header_err)}")
                                # Fallback to simple text insertion
                                try:
                                    new_page.insert_text((header_rect.x0, header_rect.y0 + 14), 
                                                          f"Translation - Page {page_num}", 
                                                          fontsize=14, 
                                                          color=(0, 0, 0.8))
                                except Exception as e:
                                    logger.error(f"Simple header insertion also failed: {str(e)}")
                            
                            # Copy content from the translated page
                            new_page.show_pdf_page(
                                fitz.Rect(0, 60, new_page.rect.width, new_page.rect.height),
                                translated_doc,
                                t_page.number
                            )
                    
                    # Clean up temporary files
                    os.unlink(translated_pdf_path)
                    
                except Exception as page_err:
                    logger.error(f"Error processing page {page_num} with markdown: {str(page_err)}")
                    # Add an empty translation page with error message
                    err_page = new_doc.new_page(width=orig_doc[page_idx].rect.width, 
                                        height=orig_doc[page_idx].rect.height)
                    err_page.insert_textbox(
                        fitz.Rect(50, 50, err_page.rect.width - 50, 100),
                        f"Error translating page {page_num}",
                        fontsize=12,
                        color=(1, 0, 0),  # Red text
                        fontname="helv",
                        align=fitz.TEXT_ALIGN_CENTER
                    )
                
                # Clean up page PDF
                os.unlink(page_pdf_path)
            
            # Save the document
            new_doc.save(output_path)
            new_doc.close()
            orig_doc.close()
            
            # Read the file back
            with open(output_path, 'rb') as f:
                pdf_data = f.read()
                
        except Exception as e:
            logger.error(f"Error in bilingual markdown PDF rebuilding: {str(e)}")
            # Fallback to standard bilingual mode
            logger.warning("Falling back to standard bilingual mode")
            return self._rebuild_pdf_bilingual(components, source_pdf_path)
        
        # Clean up the temporary file
        try:
            os.unlink(output_path)
        except:
            pass
        
        return DocumentOutput(pdf_data, 'pdf')
    
    def _draw_table_on_image(self, draw, component, rect, cjk_font_path, default_font_path, font_size=24):
        """Draw a table on the image with improved layout."""
        x0, y0, x1, y1 = rect
        
        # Draw table border
        draw.rectangle([x0, y0, x1, y1], outline=(0, 0, 0), width=3)
        
        # Calculate cell dimensions
        if not component.rows:
            return
            
        rows = len(component.rows)
        cols = len(component.rows[0]) if rows > 0 else 0
        
        if rows <= 0 or cols <= 0:
            return
            
        cell_width = (x1 - x0) / cols
        cell_height = (y1 - y0) / rows
        
        # Draw header row with different background color
        header_color = (230, 230, 230)  # Light gray
        draw.rectangle([x0, y0, x1, y0 + cell_height], fill=header_color, outline=(0, 0, 0), width=2)
        
        # Draw cells
        for i, row in enumerate(component.rows):
            for j, cell in enumerate(row):
                # Calculate cell position
                cell_x0 = x0 + j * cell_width
                cell_y0 = y0 + i * cell_height
                cell_x1 = x0 + (j + 1) * cell_width
                cell_y1 = y0 + (i + 1) * cell_height
                
                # Draw cell border
                draw.rectangle([cell_x0, cell_y0, cell_x1, cell_y1], outline=(0, 0, 0), width=1)
                
                # Add cell text
                try:
                    has_non_latin = any(ord(char) > 255 for char in cell)
                    
                    # Choose appropriate font
                    font_path = cjk_font_path if (has_non_latin and cjk_font_path) else default_font_path
                    font = self._get_font(font_path, font_size)
                    
                    # Get text dimensions
                    text_width, text_height = self._get_text_dimensions(cell, font)
                    
                    # Get cell interior dimensions (with padding)
                    padding = 4
                    interior_width = cell_width - (2 * padding)
                    interior_height = cell_height - (2 * padding)
                    
                    # Check if text needs to be wrapped
                    if text_width > interior_width:
                        # Text needs wrapping
                        words = cell.split()
                        lines = []
                        current_line = []
                        current_width = 0
                        
                        for word in words:
                            word_width, _ = self._get_text_dimensions(word + " ", font)
                            if current_width + word_width <= interior_width:
                                current_line.append(word)
                                current_width += word_width
                            else:
                                if current_line:
                                    lines.append(" ".join(current_line))
                                current_line = [word]
                                current_width = word_width
                        
                        if current_line:
                            lines.append(" ".join(current_line))
                        
                        # Draw each line of text
                        line_spacing = font_size * 1.1  # 110% of font size
                        max_lines = int(interior_height / line_spacing)
                        
                        # Truncate lines if they don't fit
                        if len(lines) > max_lines:
                            lines = lines[:max_lines-1] + ['...']
                            
                        # Calculate vertical position to center all lines
                        total_text_height = len(lines) * line_spacing
                        start_y = cell_y0 + padding + (interior_height - total_text_height) / 2
                        
                        for i, line in enumerate(lines):
                            line_y = start_y + i * line_spacing
                            # Center the text within the cell horizontally
                            line_width, _ = self._get_text_dimensions(line, font)
                            line_x = cell_x0 + padding + (interior_width - line_width) / 2
                            draw.text((line_x, line_y), line, font=font, fill=(0, 0, 0))
                    else:
                        # Center the text within the cell
                        text_x = cell_x0 + padding + (interior_width - text_width) / 2
                        text_y = cell_y0 + padding + (interior_height - text_height) / 2
                        draw.text((text_x, text_y), cell, font=font, fill=(0, 0, 0))
                        
                except Exception as e:
                    logger.error(f"Error adding cell text with PIL: {str(e)}")
                    
    def _rebuild_docx(self, components: List[DocumentComponent]) -> DocumentOutput:
        """
        Rebuild a DOCX document.
        
        Args:
            components: List of document components
            
        Returns:
            DocumentOutput object with DOCX data
        """
        # Create a new DOCX
        doc = docx.Document()
        
        # Group components by page
        page_components = {}
        for component in components:
            page_num = component.page_number
            if page_num not in page_components:
                page_components[page_num] = []
            page_components[page_num].append(component)
        
        # Process each page
        for page_num in sorted(page_components.keys()):
            # Add page break if not the first page
            if page_num > 1:
                doc.add_page_break()
                
            # Sort by vertical position
            page_components[page_num].sort(key=lambda c: getattr(c, 'position', {}).get('y0', 0))
            
            for component in page_components[page_num]:
                if isinstance(component, TextComponent):
                    # Add text
                    text = component.text
                    if not text.strip():
                        continue  # Skip empty text
                    
                    # Determine if this might be a heading based on font size
                    font_size = component.font_info.get('size', 11)
                    is_bold = component.font_info.get('bold', False)
                    
                    if font_size >= 16 or is_bold:
                        # Add as heading
                        heading_level = 1 if font_size >= 16 else 2
                        doc.add_heading(text, level=heading_level)
                    else:
                        # Add as paragraph
                        doc.add_paragraph(text)
                        
                elif isinstance(component, TableComponent):
                    # Add table
                    if not component.rows:
                        continue
                        
                    table = doc.add_table(rows=len(component.rows), cols=len(component.rows[0]))
                    
                    # Populate cells
                    for i, row in enumerate(component.rows):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = cell
                    
                    # Add space after table
                    doc.add_paragraph()
                    
                elif isinstance(component, ImageComponent):
                    # Skip images for now
                    continue
                    
                elif isinstance(component, ChartComponent):
                    # Skip charts for now
                    continue
        
        # Save DOCX to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return document output
        return DocumentOutput(buffer.getvalue(), 'docx')
