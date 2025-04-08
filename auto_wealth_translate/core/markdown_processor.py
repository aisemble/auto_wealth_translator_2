"""
Markdown-based document processing module.
"""
import fitz
import re
import os
import base64
from typing import List, Dict, Any, Optional, Tuple
import logging
from pathlib import Path
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import subprocess
import os

logger = logging.getLogger(__name__)

class MarkdownProcessor:
    """Process documents through Markdown format."""
    
    def __init__(self):
        """Initialize the processor."""
        self.md_converter = markdown.Markdown(extensions=['tables', 'fenced_code'])
        self.image_dir = None
        
    def pdf_to_markdown(self, pdf_path: str) -> str:
        """
        Convert PDF to Markdown while preserving structure.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            str: Markdown content
        """
        try:
            doc = fitz.open(pdf_path)
            md_content = []
            
            logger.info(f"Converting PDF to Markdown: {pdf_path}")
            logger.info(f"PDF has {len(doc)} pages")
            
            # Create temporary directory for images
            self.image_dir = Path(tempfile.gettempdir()) / f"md_images_{Path(pdf_path).stem}"
            self.image_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created image directory: {self.image_dir}")
            
            for page_num, page in enumerate(doc):
                # Extract text blocks with their properties
                text_dict = page.get_text("dict")
                blocks = text_dict["blocks"]
                
                logger.info(f"Processing page {page_num+1} with {len(blocks)} blocks")
                
                # Log entire page text for verification
                page_text = page.get_text()
                logger.info(f"Page {page_num+1} raw text sample: {page_text[:200]}...")
                
                # Extract images from the page
                image_list = page.get_images(full=True)
                logger.info(f"Found {len(image_list)} images on page {page_num+1}")
                
                # Process and save images
                for img_idx, img_info in enumerate(image_list):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Save image to temporary directory
                        img_filename = f"page{page_num+1}_img{img_idx}.{image_ext}"
                        img_path = self.image_dir / img_filename
                        with open(img_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        
                        # Add image reference to markdown content
                        md_content.append(f"\n![Image {page_num+1}-{img_idx}]({img_path})\n")
                        logger.info(f"Added image reference to {img_path}")
                    except Exception as img_err:
                        logger.error(f"Error extracting image {img_idx} on page {page_num+1}: {str(img_err)}")
                
                for block_idx, block in enumerate(blocks):
                    if "lines" in block:
                        for line_idx, line in enumerate(block["lines"]):
                            line_text = ""
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if not text:
                                    continue
                                    
                                # Log extracted text for debugging
                                logger.debug(f"Block {block_idx}, Line {line_idx}, Text: {text[:50]}{'...' if len(text) > 50 else ''}")
                                line_text += text + " "
                                    
                                # Determine text style
                                font_size = span["size"]
                                font_name = span["font"]
                                is_bold = "bold" in font_name.lower()
                                
                                # Add text with appropriate markdown formatting
                                if font_size > 14:
                                    md_content.append(f"# {text}")
                                elif font_size > 12:
                                    md_content.append(f"## {text}")
                                elif is_bold:
                                    md_content.append(f"**{text}**")
                                else:
                                    md_content.append(text)
                            
                            # Log full line text
                            if line_text.strip():
                                logger.debug(f"Full line text: {line_text.strip()}")
                                    
                # Add page break
                md_content.append("\n---\n")
            
            result = "\n".join(md_content)
            # Print a sample of the extracted content
            logger.info(f"Extracted markdown sample (first 500 chars): {result[:500]}...")
            logger.info(f"Total markdown content length: {len(result)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Error converting PDF to Markdown: {str(e)}")
            raise
            
    def translate_markdown(self, md_content: str, translator) -> str:
        """
        Translate Markdown content while preserving formatting.
        
        Args:
            md_content: Original Markdown content
            translator: Translation service instance
            
        Returns:
            str: Translated Markdown content
        """
        try:
            # Print a sample of the original markdown content
            logger.info(f"Original markdown sample (first 300 chars): {md_content[:300]}...")
            logger.info(f"Translating from {translator.source_lang} to {translator.target_lang}")
            
            # Split content into translatable segments
            segments = self._split_markdown(md_content)
            translated_segments = []
            
            logger.info(f"Split markdown into {len(segments)} segments")
            
            # Keep track of failed translations
            failed_segments = []
            
            for i, segment in enumerate(segments):
                if segment["type"] == "text":
                    # Skip if segment is empty or whitespace
                    if not segment["content"].strip():
                        translated_segments.append(segment)
                        continue
                    
                    # Log the original text segment
                    logger.info(f"Segment {i} (original): {segment['content'][:100]}{'...' if len(segment['content']) > 100 else ''}")
                    
                    # Check if the segment contains mathematical formulas or special characters
                    content = segment['content']
                    has_math = self._contains_math_or_special_chars(content)
                    
                    try:
                        # Get the target language code
                        target_lang = translator.target_lang
                        
                        # For Chinese translation, use a custom approach
                        if target_lang == "zh":
                            # If the segment contains math formulas, preserve it as is
                            if has_math:
                                logger.info(f"Segment {i} contains mathematical notation, preserving it as is")
                                translated_segments.append(segment)
                                continue
                                
                            logger.info(f"Using specialized approach for Chinese translation of segment {i}")
                            
                            # Create explicit instruction for Chinese translation
                            instruction = f"""
                            Translate the following text from {translator.language_names.get(translator.source_lang, translator.source_lang)} to Chinese (Simplified).
                            
                            Rules:
                            1. ONLY respond with the translated text
                            2. Keep all formatting and special characters
                            3. Use appropriate Chinese financial terminology
                            4. Ensure output is in UTF-8 encoded Chinese characters
                            5. Do not add any comments, explanations, or notes
                            6. DO NOT translate mathematical formulas, variable names, or equations - keep them exactly as they are
                            
                            Text to translate:
                            {segment['content']}
                            """
                            
                            # Use direct call with temperature = 0 for consistency
                            translated_text = translator._translate_with_openai(
                                text=instruction,
                                target_lang="zh",
                                temperature=0.0
                            )
                            
                            # Verify Chinese characters are present
                            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in translated_text)
                            if not has_chinese:
                                logger.warning(f"No Chinese characters found in translation output for segment {i}! Content: {translated_text[:100]}")
                                
                                # For segments with special characters or math formulas, preserve original
                                if self._contains_complex_notation(segment['content']):
                                    logger.info(f"Segment {i} contains complex notation, preserving original")
                                    translated_segments.append(segment)
                                    continue
                                
                                # Try one more time with simpler instruction
                                retry_instruction = f"将以下文本翻译成中文(不要添加任何解释,只需给出翻译结果,保留所有数学公式和特殊符号不变):\n\n{segment['content']}"
                                translated_text = translator._translate_with_openai(
                                    text=retry_instruction, 
                                    target_lang="zh",
                                    temperature=0.0
                                )
                                
                                # Check again
                                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in translated_text)
                                if not has_chinese:
                                    logger.error(f"Second attempt also failed to produce Chinese characters for segment {i}")
                                    # Preserve original content if translation fails
                                    translated_segments.append(segment)
                                    # Record failed segment for reference
                                    failed_segments.append(i)
                                    continue
                                else:
                                    logger.info(f"Second attempt successfully produced Chinese characters for segment {i}")
                            else:
                                logger.info(f"Chinese characters verified in translation output for segment {i}")
                                
                        else:
                            # For other languages, use standard approach
                            translated_text = translator._translate_with_openai(
                                text=segment["content"],
                                target_lang=target_lang,
                                temperature=0.3
                            )
                        
                        # Log the translated segment
                        logger.info(f"Segment {i} (translated): {translated_text[:100]}{'...' if len(translated_text) > 100 else ''}")
                        
                        # Store the translated segment
                        translated_segments.append({
                            "type": "text",
                            "content": translated_text
                        })
                        
                    except Exception as e:
                        logger.error(f"Error translating segment {i}: {str(e)}")
                        # Fall back to original content on error
                        translated_segments.append(segment)
                        failed_segments.append(i)
                else:
                    # Keep formatting intact
                    translated_segments.append(segment)
            
            # Reconstruct markdown
            result = self._reconstruct_markdown(translated_segments)
            
            # Print a sample of the translated content
            logger.info(f"Translated markdown sample (first 300 chars): {result[:300]}...")
            logger.info(f"Total translated content length: {len(result)} characters")
            
            # Check if we have any Chinese characters in result when target is Chinese
            if translator.target_lang == "zh":
                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in result)
                if not has_chinese:
                    logger.error("No Chinese characters found in the final translated content!")
                else:
                    logger.info("Chinese characters verified in the translated content")
            
            # Post-process the translated content if there are failed segments or if it's Chinese
            # This step improves the formatting and flow of the translated document
            if failed_segments or translator.target_lang == "zh":
                result = self._post_process_translation(result, translator, failed_segments)
                logger.info("Applied post-processing to improve translation quality and formatting")
            
            return result
            
        except Exception as e:
            logger.error(f"Error translating Markdown: {str(e)}")
            raise
            
    def _split_markdown(self, md_content: str) -> List[Dict[str, Any]]:
        """Split markdown into translatable segments."""
        segments = []
        current_text = []
        
        for line in md_content.split("\n"):
            # Check for mathematical formulas and preserve them
            if (line.strip().startswith('$') and line.strip().endswith('$')) or \
               ('\\begin{' in line and '\\end{' in line) or \
               (self._contains_math_or_special_chars(line)):
                # Save any current text before the formula
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Check for image references (don't translate these)
            if line.strip().startswith("![") and "](" in line and line.endswith(")"):
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Check for headers
            if line.startswith("#"):
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Check for code blocks
            if line.strip().startswith("```") or line.strip().startswith("~~~"):
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Check for formatting
            if line.startswith(("**", "*", "`", ">", "-", "1.", "|", "---")) or line.strip() == "":
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Check for special characters that might indicate formulas or technical content
            if self._contains_complex_notation(line):
                if current_text:
                    segments.append({
                        "type": "text",
                        "content": "\n".join(current_text)
                    })
                    current_text = []
                segments.append({
                    "type": "format",
                    "content": line
                })
                continue
                
            # Regular text
            current_text.append(line)
            
        if current_text:
            segments.append({
                "type": "text",
                "content": "\n".join(current_text)
            })
            
        return segments
        
    def _reconstruct_markdown(self, segments: List[Dict[str, Any]]) -> str:
        """Reconstruct markdown from translated segments."""
        return "\n".join(segment["content"] for segment in segments)
        
    def markdown_to_docx(self, md_content: str, output_path: str):
        """
        Convert Markdown to DOCX.
        
        Args:
            md_content: Markdown content
            output_path: Output DOCX path
        """
        try:
            # Convert markdown to HTML
            html = self.md_converter.convert(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            
            # Create new document
            doc = Document()
            
            # Process HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li', 'table', 'img']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    p = doc.add_paragraph()
                    p.style = f'Heading {level}'
                    p.add_run(element.get_text())
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    p.add_run(element.get_text())
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        p = doc.add_paragraph()
                        p.style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        p.add_run(li.get_text())
                elif element.name == 'table':
                    table = doc.add_table(rows=1, cols=1)
                    for row in element.find_all('tr'):
                        cells = row.find_all(['td', 'th'])
                        if len(cells) > table.columns:
                            table.add_column()
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            row_cells[i].text = cell.get_text()
                elif element.name == 'img':
                    # Handle images
                    src = element.get('src')
                    if src and os.path.exists(src):
                        try:
                            p = doc.add_paragraph()
                            r = p.add_run()
                            r.add_picture(src, width=Inches(6))  # Adjust width as needed
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            logger.info(f"Added image from {src} to DOCX document")
                        except Exception as img_err:
                            logger.error(f"Error adding image to DOCX: {str(img_err)}")
            
            # Save document
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error converting Markdown to DOCX: {str(e)}")
            raise
            
    def markdown_to_pdf(self, markdown_content, output_path):
        """
        Convert markdown to PDF.
        
        Args:
            markdown_content: Markdown content
            output_path: Path to output PDF file
        """
        logger.info(f"Converting Markdown to PDF: {output_path}")
        
        try:
            # First try WeasyPrint if available
            try:
                from weasyprint import HTML, CSS
                from weasyprint.text.fonts import FontConfiguration
            
                # Set up font configuration
                font_config = FontConfiguration()
                
                # Create CSS for better formatting
                css = CSS(string='''
                    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans&family=Noto+Sans+SC&display=swap');
                    body {
                        font-family: 'Noto Sans', 'Noto Sans SC', sans-serif;
                        margin: 3em;
                        line-height: 1.5;
                        font-size: 14px;
                    }
                    h1 {
                        font-size: 40px;
                        margin-bottom: 1em;
                        line-height: 1.2;
                    }
                    h2 {
                        font-size: 30px;
                        margin-bottom: 0.8em;
                        line-height: 1.3;
                    }
                    h3 {
                        font-size: 24px;
                        margin-bottom: 0.6em;
                        line-height: 1.4;
                    }
                    p {
                        font-size: 14px;
                        margin-bottom: 1em;
                    }
                    img {
                        max-width: 100%;
                        margin: 1em 0;
                    }
                    table {
                        border-collapse: collapse;
                        margin: 1em 0;
                        width: 100%;
                    }
                    th, td {
                        border: 1px solid #ddd;
                        padding: 0.5em;
                        font-size: 14px;
                    }
                    th {
                        background-color: #f5f5f5;
                        font-weight: bold;
                    }
                    code {
                        font-size: 13px;
                        background-color: #f5f5f5;
                        padding: 0.2em 0.4em;
                        border-radius: 3px;
                    }
                    pre {
                        background-color: #f5f5f5;
                        padding: 1em;
                        border-radius: 5px;
                        overflow-x: auto;
                    }
                    pre code {
                        font-size: 13px;
                        background-color: transparent;
                        padding: 0;
                    }
                    ul, ol {
                        font-size: 14px;
                        margin-bottom: 1em;
                        padding-left: 1.5em;
                    }
                    li {
                        margin-bottom: 0.5em;
                    }
                    blockquote {
                        font-size: 14px;
                        border-left: 4px solid #ddd;
                        margin: 1em 0;
                        padding-left: 1em;
                        color: #666;
                    }
                ''', font_config=font_config)
                
                # Convert to HTML and then to PDF
                html = markdown.markdown(
                    markdown_content,
                    extensions=['tables', 'fenced_code']
                )
                HTML(string=html).write_pdf(output_path, stylesheets=[css], font_config=font_config)
                logger.info("Converted to PDF using WeasyPrint")
                return
            except ImportError:
                logger.warning("WeasyPrint not available. Trying alternative methods...")
            
            # Next try PyMuPDF
            try:
                logger.info("Trying PyMuPDF for PDF generation")
                import fitz
                
                # Create a new PDF document
                doc = fitz.open()
                page = doc.new_page()
                
                # Convert markdown to HTML for better rendering
                html = markdown.markdown(
                    markdown_content,
                    extensions=['tables', 'fenced_code']
                )
                
                # Extract and add images first
                img_pattern = r'!\[(.*?)\]\((.*?)\)'
                img_matches = re.findall(img_pattern, markdown_content)
                
                y_position = 50  # Starting y position for content
                
                for _, img_path in img_matches:
                    if os.path.exists(img_path):
                        try:
                            img_rect = fitz.Rect(50, y_position, 550, y_position + 200)
                            page.insert_image(img_rect, filename=img_path)
                            y_position += 220  # Move down for next content
                            logger.info(f"Added image from {img_path} to PDF")
                        except Exception as e:
                            logger.error(f"Error adding image: {str(e)}")
                
                # Add text content
                # Strip markdown image syntax to avoid duplication
                text_content = re.sub(img_pattern, '', markdown_content)
                
                # Clean up other markdown syntax
                text_content = re.sub(r'#{1,6}\s+(.*)', r'\1', text_content)  # Headers
                text_content = re.sub(r'\*\*(.*?)\*\*', r'\1', text_content)  # Bold
                text_content = re.sub(r'\*(.*?)\*', r'\1', text_content)      # Italic
                
                # Add the text - note we're using insert_text to avoid the fontname parameter issue
                try:
                    rect = fitz.Rect(50, y_position, 550, 800)
                    page.insert_text(rect.tl, text_content, fontsize=11, color=(0,0,0))
                except Exception as e:
                    logger.error(f"PyMuPDF error: {str(e)}")
                    # Try a simpler approach with just text
                    try:
                        page.insert_text((50, y_position), "Translated content:", fontsize=12)
                        lines = text_content.split('\n')
                        for i, line in enumerate(lines[:40]):  # Limit to first 40 lines
                            page.insert_text((50, y_position + 20 + i*15), line, fontsize=10)
                    except Exception as e2:
                        logger.error(f"Second PyMuPDF error: {str(e2)}")
                        raise
                
                # Save the document
                doc.save(output_path)
                doc.close()
                logger.info("Converted to PDF using PyMuPDF")
                return
            except Exception as e:
                logger.error(f"PyMuPDF error: {str(e)}")
            
            # Try pandoc as last resort
            try:
                logger.info("Trying pandoc for PDF generation")
                
                # Save markdown to temporary file
                with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as temp_md:
                    temp_md_path = temp_md.name
                    temp_md.write(markdown_content.encode('utf-8'))
                
                # Build pandoc command - use xelatex for CJK support
                cmd = f"pandoc {temp_md_path} -o {output_path} --pdf-engine=xelatex -V mainfont=Noto Sans -V CJKmainfont=Noto Sans CJK SC"
                
                logger.info(f"Running pandoc command: {cmd}")
                result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
                
                # Clean up temp file
                os.unlink(temp_md_path)
                
                if result.returncode != 0:
                    logger.error(f"Pandoc error: {result.stderr}")
                    raise Exception(f"Pandoc failed: {result.stderr}")
                
                logger.info("Converted to PDF using Pandoc")
                return
            except Exception as e:
                logger.error(f"Pandoc error: {str(e)}")
            
            # If all else fails, create a basic PDF with reportlab
            try:
                logger.info("Trying ReportLab for basic PDF generation")
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.utils import simpleSplit
                
                # Create a canvas
                c = canvas.Canvas(output_path, pagesize=letter)
                width, height = letter
                
                # Try to register a font that supports CJK
                try:
                    # Try to find a suitable font
                    font_path = None
                    # Common paths by OS
                    font_dirs = []
                    if os.name == 'posix':  # macOS, Linux
                        font_dirs = [
                            "/System/Library/Fonts",
                            "/Library/Fonts",
                            "/usr/share/fonts"
                        ]
                    elif os.name == 'nt':  # Windows
                        font_dirs = [
                            "C:\\Windows\\Fonts"
                        ]
                    
                    # Look for a CJK font
                    for font_dir in font_dirs:
                        if os.path.exists(font_dir):
                            for font_file in os.listdir(font_dir):
                                if font_file.lower().endswith(('.ttf', '.ttc')) and any(x in font_file.lower() for x in ['cjk', 'chinese', 'japanese', 'korean', 'ming', 'gothic', 'hei', 'sans']):
                                    font_path = os.path.join(font_dir, font_file)
                                    break
                            if font_path:
                                break
                    
                    if font_path:
                        # Register the font with ReportLab
                        pdfmetrics.registerFont(TTFont('CJKFont', font_path))
                        font_name = 'CJKFont'
                    else:
                        font_name = 'Helvetica'
                except Exception:
                    font_name = 'Helvetica'  # Default to Helvetica if font registration fails
                
                # Write the content
                y = height - 40
                c.setFont(font_name, 12)
                c.drawString(30, y, "Translated Document")
                y -= 20
                
                c.setFont(font_name, 10)
                
                # Break content into pages
                lines = []
                for para in markdown_content.split('\n\n'):
                    # Clean up markdown syntax
                    para = re.sub(r'#{1,6}\s+(.*)', r'\1', para)  # Headers
                    para = re.sub(r'\*\*(.*?)\*\*', r'\1', para)  # Bold
                    para = re.sub(r'\*(.*?)\*', r'\1', para)      # Italic
                    para = re.sub(r'!\[(.*?)\]\((.*?)\)', '[Image]', para)  # Images
                    
                    # Split paragraph into lines that fit on the page
                    para_lines = simpleSplit(para, font_name, 10, width - 60)
                    lines.extend(para_lines)
                    lines.append('')  # Add blank line between paragraphs
                
                # Write lines to pages
                line_height = 14
                for line in lines:
                    if y < 40:  # Start a new page if near bottom
                        c.showPage()
                        y = height - 40
                        c.setFont(font_name, 10)
                    
                    try:
                        c.drawString(30, y, line)
                    except:
                        # Fall back to basic ASCII if there are encoding issues
                        c.drawString(30, y, line.encode('ascii', 'replace').decode())
                    
                    y -= line_height
                
                # Save the PDF
                c.save()
                logger.info("Created basic PDF using ReportLab")
                return
            except Exception as e:
                logger.error(f"ReportLab error: {str(e)}")
            
            # Ultimate fallback - create a simple text file with .pdf extension
            logger.warning("All PDF generation methods failed. Creating text file with .pdf extension")
            with open(output_path, 'wb') as f:
                f.write(b"%PDF-1.4\n")  # PDF header
                f.write(b"1 0 obj\n<</Type /Catalog /Pages 2 0 R>>\nendobj\n")
                f.write(b"2 0 obj\n<</Type /Pages /Kids [3 0 R] /Count 1>>\nendobj\n")
                f.write(b"3 0 obj\n<</Type /Page /Parent 2 0 R /Resources 4 0 R /MediaBox [0 0 612 792] /Contents 5 0 R>>\nendobj\n")
                f.write(b"4 0 obj\n<</Font <</F1 <<\n/Type /Font /Subtype /Type1 /BaseFont /Helvetica>>>>>>\nendobj\n")
                
                # Escape special characters and convert to PDF content stream format
                content = "Translation completed but PDF generation failed. Please check the logs."
                pdf_content = f"BT\n/F1 12 Tf\n50 700 Td\n({content}) Tj\nET"
                
                f.write(f"5 0 obj\n<</Length {len(pdf_content)}>>\nstream\n{pdf_content}\nendstream\nendobj\n".encode())
                f.write(b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000056 00000 n \n0000000111 00000 n \n0000000212 00000 n \n0000000293 00000 n \n")
                f.write(b"trailer\n<</Size 6 /Root 1 0 R>>\nstartxref\n394\n%%EOF\n")
            
            logger.warning(f"Created minimal PDF placeholder at {output_path}")
            
        except Exception as e:
            logger.error(f"Error converting Markdown to PDF: {str(e)}")
            raise 

    def _contains_math_or_special_chars(self, text):
        """Check if text contains mathematical notation or special characters."""
        # Check for common math symbols
        math_symbols = ['∑', '∫', '∂', '∇', '∈', '∉', '∋', '∝', '∞', '∠', '∥', '∟', '∩', '∪', '∧', '∨', '≠', '≡', '≤', '≥', '⊂', '⊃', '⊆', '⊇']
        if any(symbol in text for symbol in math_symbols):
            return True
        
        # Check for variable names with subscripts or superscripts
        if re.search(r'[a-zA-Z]_[a-zA-Z0-9]', text) or re.search(r'[a-zA-Z]\^[a-zA-Z0-9]', text):
            return True
        
        # Check for LaTeX-like formulas
        if re.search(r'\$.*\$', text) or re.search(r'\\\(.*\\\)', text) or re.search(r'\\\[.*\\\]', text):
            return True
        
        # Check for fractions and equations
        if '/' in text and re.search(r'[a-zA-Z0-9]/[a-zA-Z0-9]', text):
            return True
        
        # Check for Greek letters
        if re.search(r'\\alpha|\\beta|\\gamma|\\delta|\\epsilon|\\zeta|\\eta|\\theta|\\lambda|\\mu|\\nu|\\xi|\\pi|\\rho|\\sigma|\\tau|\\upsilon|\\phi|\\chi|\\psi|\\omega', text):
            return True
        
        return False

    def _contains_complex_notation(self, text):
        """Check if text contains complex mathematical notation or special symbols that shouldn't be translated."""
        # Check for special unicode characters that often cause issues
        if any(ord(char) > 127 and ord(char) not in range(0x4e00, 0x9fff) for char in text):
            return True
        
        # Check for mathematical expressions
        if re.search(r'[−×÷=≤≥≈∑∏∞∫∂∇]', text):
            return True
        
        # Check for technical notations like variable names
        if re.search(r'[a-zA-Z][0-9]', text) or re.search(r'[a-zA-Z]_[a-zA-Z0-9]', text) or re.search(r'[a-zA-Z]\^[a-zA-Z0-9]', text):
            return True
        
        # Check for matrices, vectors, or tensors (often in scientific papers)
        if re.search(r'[𝒙𝑩𝑨𝑻𝑴𝑽𝑾]', text) or re.search(r'[a-zA-Z]←[a-zA-Z]', text):
            return True
        
        # Check for inline math expressions with parentheses, brackets
        if re.search(r'\([a-z0-9]+[+\-*/][a-z0-9]+\)', text):
            return True
        
        # Check for superscripts and subscripts
        if re.search(r'[⁰¹²³⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉]', text):
            return True
        
        # Check for special mathematical symbols
        math_symbols = ['∑', '∫', '∂', '∇', '∈', '∉', '∋', '∝', '∞', '∠', '∥', '∟', '∩', '∪', '∧', '∨', '≠', '≡', '≤', '≥', '⊂', '⊃', '⊆', '⊇', '±', '·', '×', '÷', '→', '⇒', '⇔']
        if any(symbol in text for symbol in math_symbols):
            return True
        
        # Check for LaTeX-like commands
        if re.search(r'\\[a-zA-Z]+', text):
            return True
        
        # A segment with high percentage of special characters or numbers is likely a formula
        total_chars = len(text)
        special_chars = sum(1 for char in text if not char.isalpha() or char in "()[]{}.,;:+-*/=<>")
        if total_chars > 0 and special_chars / total_chars > 0.35:  # If more than 35% are special chars
            return True
        
        # Check for consecutive symbols that likely indicate equations
        if re.search(r'[+\-*/=<>]{2,}', text):
            return True
        
        return False

    def _post_process_translation(self, translated_content: str, translator, failed_segments: list = None) -> str:
        """
        Post-process the translated markdown to improve formatting and flow.
        
        Args:
            translated_content: The translated markdown content
            translator: Translation service instance
            failed_segments: List of indices of segments that failed to translate
            
        Returns:
            str: Improved translated markdown content
        """
        logger.info("Post-processing translated content to improve formatting and flow")
        
        # If the content is too long, process it in chunks
        if len(translated_content) > 4000:
            chunks = []
            content_lines = translated_content.split('\n')
            current_chunk = []
            current_length = 0
            
            for line in content_lines:
                line_length = len(line)
                if current_length + line_length > 4000:
                    chunks.append('\n'.join(current_chunk))
                    current_chunk = [line]
                    current_length = line_length
                else:
                    current_chunk.append(line)
                    current_length += line_length
                
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
            
            # Process each chunk
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                logger.info(f"Post-processing chunk {i+1}/{len(chunks)}")
                processed_chunk = self._process_content_chunk(chunk, translator)
                processed_chunks.append(processed_chunk)
            
            return '\n'.join(processed_chunks)
        else:
            # Process the entire content at once
            return self._process_content_chunk(translated_content, translator)
        
    def _process_content_chunk(self, content: str, translator) -> str:
        """Process a chunk of translated content to improve formatting and coherence."""
        target_lang = translator.target_lang
        source_lang = translator.source_lang
        
        # Create instructions for post-processing
        if target_lang == "zh":
            instruction = f"""
            You are a professional editor for translated markdown documents. The following text has been 
            machine-translated from {translator.language_names.get(source_lang, source_lang)} to Chinese.
            
            Your task is to improve the formatting and coherence of the text while preserving the original 
            meaning and content structure.
            
            Instructions:
            1. Fix any awkward translations or incomplete sentences
            2. Maintain the original markdown formatting (headings, lists, code blocks, etc.)
            3. Ensure proper Chinese grammar and phrasing
            4. DO NOT translate or modify any technical terms, code, or mathematical formulas
            5. Preserve all image references and links exactly as they appear
            6. Make sure tables remain properly formatted in markdown
            7. ONLY return the improved markdown document without any additional comments
            
            Here's the content to improve:
            
            {content}
            """
        else:
            instruction = f"""
            You are a professional editor for translated markdown documents. The following text has been 
            machine-translated from {translator.language_names.get(source_lang, source_lang)} to {translator.language_names.get(target_lang, target_lang)}.
            
            Your task is to improve the formatting and coherence of the text while preserving the original 
            meaning and content structure.
            
            Instructions:
            1. Fix any awkward translations or incomplete sentences
            2. Maintain the original markdown formatting (headings, lists, code blocks, etc.)
            3. Ensure proper grammar and phrasing
            4. DO NOT translate or modify any technical terms, code, or mathematical formulas
            5. Preserve all image references and links exactly as they appear
            6. Make sure tables remain properly formatted in markdown
            7. ONLY return the improved markdown document without any additional comments
            
            Here's the content to improve:
            
            {content}
            """
        
        try:
            # Use a slightly higher temperature for more creative reformatting
            improved_content = translator._translate_with_openai(
                text=instruction,
                target_lang=target_lang,
                temperature=0.4
            )
            
            # Verify the improved content still contains the target language if it's Chinese
            if target_lang == "zh":
                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in improved_content)
                if not has_chinese:
                    logger.warning("Post-processing removed Chinese characters, reverting to original translation")
                    return content
                
            return improved_content
        except Exception as e:
            logger.error(f"Error in post-processing: {str(e)}")
            # If post-processing fails, return the original translation
            return content 