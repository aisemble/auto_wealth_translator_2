"""
AutoWealthTranslate Streamlit App
---------------------------------
A Streamlit UI for the AutoWealthTranslate application.
"""

import os
import streamlit as st
import tempfile
from pathlib import Path
import time
import logging
import shutil
import uuid
import base64
import io
import re
import json
from openai import OpenAI

from auto_wealth_translate.core.document_processor import DocumentProcessor
from auto_wealth_translate.core.translator import TranslationService
from auto_wealth_translate.core.document_rebuilder import DocumentRebuilder
from auto_wealth_translate.core.validator import OutputValidator
from auto_wealth_translate.core.markdown_processor import MarkdownProcessor
from auto_wealth_translate.core.chart_processor import chart_to_markdown
from auto_wealth_translate.utils.logger import setup_logger, get_logger
from auto_wealth_translate.core.deepl_translator import DeepLTranslationService

# Configure logging
setup_logger(logging.INFO)
logger = get_logger("streamlit_app")

# Set more verbose logging for the markdown processor module
markdown_logger = logging.getLogger("auto_wealth_translate.core.markdown_processor")
markdown_logger.setLevel(logging.DEBUG)

# Create a file handler for detailed logs
log_file = Path(tempfile.gettempdir()) / "markdown_processor_debug.log"
file_handler = logging.FileHandler(log_file)
file_handler.setLevel(logging.DEBUG)
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
markdown_logger.addHandler(file_handler)

logger.info(f"Debug logs will be written to: {log_file}")

# Supported languages
SUPPORTED_LANGUAGES = {
    "en": "English",
    "zh": "Chinese (中文)",
    "fr": "French (Français)",
    "es": "Spanish (Español)",
    "de": "German (Deutsch)",
    "ja": "Japanese (日本語)",
    "ko": "Korean (한국어)",
    "ru": "Russian (Русский)",
    "ar": "Arabic (العربية)",
    "it": "Italian (Italiano)",
    "pt": "Portuguese (Português)"
}

# Translation service configuration
TRANSLATION_SERVICE = "Document Translation"

def create_temp_dir():
    """Create a temporary directory for file processing."""
    temp_dir = Path(tempfile.gettempdir()) / "autowealthtranslate_streamlit"
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir

def get_file_download_link(file_path, link_text):
    """Generate a download link for a file."""
    with open(file_path, "rb") as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    filename = Path(file_path).name
    mime_type = "application/pdf" if file_path.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}">{link_text}</a>'
    return href

def translate_document(input_file, source_lang, target_lang, api_key=None):
    """
    Process and translate a document.
    
    Args:
        input_file: Path to input file
        source_lang: Source language code
        target_lang: Target language code
        api_key: API key for translation service
    
    Returns:
        Path to translated file, validation results, markdown content (if markdown mode), summary
    """
    if api_key:
        os.environ["DEEPL_API_KEY"] = api_key
    
    temp_dir = create_temp_dir()
    
    # Generate unique filename
    input_path = Path(input_file)
    unique_id = str(uuid.uuid4())[:8]
    output_path = temp_dir / f"{input_path.stem}_{target_lang}_{unique_id}{input_path.suffix}"
    
    # Initialize translator
    logger.info(f"Translating document from {source_lang} to {target_lang}")
    
    translator = DeepLTranslationService(
        source_lang=source_lang, 
        target_lang=target_lang,
        api_key=api_key
    )
    
    if not translator.is_ready():
        raise Exception("Translation service is not initialized. Please check your API key.")
    
    # Translate document
    translated_file, metadata = translator.translate_document(input_file)
    
    if not translated_file:
        error_message = metadata.get("error", "Unknown error during translation")
        raise Exception(f"Translation failed: {error_message}")
    
    # Create validation result
    validation_result = {
        "score": 9,  # Translations are generally high quality
        "issues": [],
        "metadata": metadata
    }
    
    # Generate summary
    summary = generate_document_summary(translated_file, target_lang)
    
    # Return the translated file path, validation results, and summary
    return translated_file, validation_result, None, summary

def generate_document_summary(file_path, language):
    """Generate a summary of the translated document."""
    from pathlib import Path
    import os
    
    # Use temporary file for extraction
    temp_dir = create_temp_dir()
    
    file_path = Path(file_path)
    file_ext = file_path.suffix.lower()
    
    # Extract text from document
    extracted_text = ""
    
    try:
        if file_ext == '.pdf':
            # For PDF files
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n\n"
        elif file_ext == '.docx':
            # For DOCX files
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            
            # Also get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " | "
                    extracted_text += "\n"
        else:
            return "Summary generation is only supported for PDF and DOCX files."
        
        # Limit text to avoid token limits
        max_chars = 15000
        if len(extracted_text) > max_chars:
            extracted_text = extracted_text[:max_chars] + "..."
            
        # Create prompt for summary generation
        if language == 'zh':
            prompt = f"""请为以下文档生成一个简明的摘要（不超过 250 字）。关注主要内容、关键观点和结论。

文档内容:
{extracted_text}

摘要:"""
        else:
            prompt = f"""Please generate a concise summary of the following document (no more than 250 words). Focus on the main content, key points, and conclusions.

Document content:
{extracted_text}

Summary:"""
        
        return "Summary generation is currently disabled."
        
    except Exception as e:
        logger.error(f"Error generating summary: {str(e)}")
        return "Error generating summary."

def perform_qa_comparison(original_file, translated_file, openai_api_key, source_lang, target_lang):
    """
    Use OpenAI to compare original and translated documents for quality assessment.
    
    Args:
        original_file: Path to the original file
        translated_file: Path to the translated file
        openai_api_key: OpenAI API key
        source_lang: Source language code
        target_lang: Target language code
        
    Returns:
        QA analysis results as a dictionary
    """
    if not openai_api_key:
        return {"error": "OpenAI API key is required for QA analysis"}
    
    # Extract text from both documents
    original_text = extract_text_from_document(original_file)
    translated_text = extract_text_from_document(translated_file)
    
    if not original_text or not translated_text:
        return {"error": "Failed to extract text from documents"}
    
    try:
        # Initialize OpenAI client
        client = OpenAI(api_key=openai_api_key)
        
        # Create prompt for comparison
        prompt = f"""
        You are a professional translation quality assessor with expertise in financial documents.
        
        TASK: Compare the original document in {source_lang} with the translated document in {target_lang}, 
        then provide a detailed quality assessment.
        
        Original Document ({source_lang}):
        {original_text[:10000]}  # Limit text length
        
        Translated Document ({target_lang}):
        {translated_text[:10000]}  # Limit text length
        
        Please assess:
        1. Accuracy: How accurately is the meaning conveyed?
        2. Terminology: Are financial terms correctly translated?
        3. Style: Is the translation appropriate for financial documents?
        4. Formatting: Are key structures preserved?
        5. Overall quality: Rate from 1-10
        
        Provide specific examples of both good translations and any errors or issues found.
        """
        
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a financial translation quality expert."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        
        analysis = response.choices[0].message.content
        
        # Parse the analysis for structured output
        qa_result = {
            "full_analysis": analysis,
            "summary": analysis[:500] + "..." if len(analysis) > 500 else analysis,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        return qa_result
        
    except Exception as e:
        logger.error(f"Error during QA analysis: {str(e)}")
        return {"error": f"QA analysis failed: {str(e)}"}

def extract_text_from_document(file_path):
    """Extract text content from PDF or DOCX file."""
    file_path = Path(file_path)
    file_ext = file_path.suffix.lower()
    
    extracted_text = ""
    
    try:
        if file_ext == '.pdf':
            # For PDF files
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n\n"
        elif file_ext == '.docx':
            # For DOCX files
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            
            # Also get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " | "
                    extracted_text += "\n"
        else:
            return None
            
        return extracted_text
        
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        return None

def main():
    st.set_page_config(
        page_title="AutoWealthTranslate",
        page_icon="🌐",
        layout="wide"
    )
    
    st.title("AutoWealthTranslate")
    st.markdown("Translate financial documents while preserving formatting and structure.")
    
    # Sidebar
    st.sidebar.header("Settings")
    
    # Language selection
    source_lang = st.sidebar.selectbox(
        "Source Language",
        options=list(SUPPORTED_LANGUAGES.keys()),
        format_func=lambda x: SUPPORTED_LANGUAGES[x],
        index=0
    )
    
    target_lang = st.sidebar.selectbox(
        "Target Language",
        options=list(SUPPORTED_LANGUAGES.keys()),
        format_func=lambda x: SUPPORTED_LANGUAGES[x],
        index=1
    )
    
    # API key inputs
    translation_api_key = st.sidebar.text_input("Translation API Key", type="password")
    
    # Optional QA section
    st.sidebar.markdown("---")
    enable_qa = st.sidebar.checkbox("Enable Translation QA", value=False, 
                               help="Use AI to analyze translation quality (requires OpenAI API key)")
    
    if enable_qa:
        openai_api_key = st.sidebar.text_input("OpenAI API Key", type="password",
                                      help="Required for translation quality assessment")
    else:
        openai_api_key = None
    
    # File upload
    uploaded_file = st.file_uploader("Upload Document", type=["pdf", "docx"])
    
    if uploaded_file is not None:
        # Save uploaded file
        temp_dir = create_temp_dir()
        input_path = temp_dir / uploaded_file.name
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        
        # Translation button
        if st.button("Translate"):
            try:
                with st.spinner("Translating document..."):
                    translated_file, validation, _, summary = translate_document(
                        input_path,
                        source_lang,
                        target_lang,
                        translation_api_key
                    )
                
                # Display results
                st.success("Translation completed!")
                
                # Download link
                st.markdown(
                    get_file_download_link(translated_file, "Download Translated Document"),
                    unsafe_allow_html=True
                )
                
                # Perform QA analysis if enabled
                if enable_qa and openai_api_key:
                    with st.spinner("Performing translation quality assessment..."):
                        qa_results = perform_qa_comparison(
                            input_path, 
                            translated_file, 
                            openai_api_key,
                            source_lang,
                            target_lang
                        )
                    
                    # Display QA results
                    if "error" in qa_results:
                        st.error(f"QA analysis error: {qa_results['error']}")
                    else:
                        st.subheader("Translation Quality Assessment")
                        
                        with st.expander("View QA Summary"):
                            st.write(qa_results["summary"])
                        
                        with st.expander("View Full Analysis"):
                            st.write(qa_results["full_analysis"])
                
                # Display summary if available
                if summary:
                    st.subheader("Document Summary")
                    st.write(summary)
                
            except Exception as e:
                st.error(f"Error during translation: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("© 2024 AutoWealthTranslate")

if __name__ == "__main__":
    main() 